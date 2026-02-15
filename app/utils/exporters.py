import os
import re
import tempfile
import logging
import html
import secrets
from datetime import datetime, timezone
import markdown
import bleach
# from weasyprint import HTML, CSS  # Commented out due to system dependencies
from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
import json

logger = logging.getLogger(__name__)

# SECURITY: Maximum document size for export (10MB)
MAX_EXPORT_CONTENT_SIZE = 10 * 1024 * 1024

# SECURITY: Maximum line count for DOCX export to prevent DoS
MAX_EXPORT_LINES = 100000

# SECURITY: Allowed HTML tags and attributes for sanitization
ALLOWED_TAGS = [
    'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
    'p', 'br', 'hr',
    'ul', 'ol', 'li',
    'blockquote', 'pre', 'code',
    'table', 'thead', 'tbody', 'tr', 'th', 'td',
    'strong', 'em', 'b', 'i', 'u', 's', 'del', 'ins',
    'a', 'img',
    'div', 'span',
    'sup', 'sub',
]

ALLOWED_ATTRIBUTES = {
    '*': ['class', 'id'],
    'a': ['href', 'title', 'rel'],
    'img': ['src', 'alt', 'title', 'width', 'height'],
    'th': ['colspan', 'rowspan'],
    'td': ['colspan', 'rowspan'],
}


def _sanitize_title_for_filename(title: str | None, max_length: int = 50) -> str:
    """Sanitize document title for safe use in filenames.

    Removes dangerous characters to prevent path traversal attacks.
    Returns 'document' as fallback for None or empty titles.
    """
    if not title:
        return 'document'
    # Remove dangerous characters (only allow alphanumeric, spaces, dots, hyphens, underscores)
    sanitized = re.sub(r'[^\w\s.-]', '', title)
    # Replace multiple spaces with single underscore
    sanitized = re.sub(r'\s+', '_', sanitized)
    # Remove leading dots to prevent hidden files
    sanitized = sanitized.lstrip('.')
    # Truncate to max length
    return sanitized[:max_length] if sanitized else 'document'


class DocumentExporter:
    """Handle document export to various formats"""

    def __init__(self, document):
        # SECURITY: Validate document content size before export
        content = document.markdown_content or ''
        content_size = len(content.encode('utf-8'))
        if content_size > MAX_EXPORT_CONTENT_SIZE:
            raise ValueError(f"Document too large for export: {content_size} bytes exceeds {MAX_EXPORT_CONTENT_SIZE} limit")

        self.document = document
        self.temp_dir = tempfile.mkdtemp()
    
    def cleanup(self):
        """Clean up temporary files"""
        import shutil
        try:
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            logger.debug("Failed to cleanup temp dir %s: %s", self.temp_dir, e)
    
    def _generate_css_styles(self) -> str:
        """Generate CSS styles for HTML export"""
        return """
            <style>
                body {
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                    line-height: 1.6;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                    color: #333;
                }
                h1, h2, h3, h4, h5, h6 {
                    color: #2c3e50;
                    margin-top: 2em;
                    margin-bottom: 1em;
                }
                h1 { border-bottom: 2px solid #3498db; padding-bottom: 10px; }
                h2 { border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }
                code {
                    background: #f8f9fa;
                    padding: 2px 6px;
                    border-radius: 3px;
                    font-family: 'Monaco', 'Consolas', monospace;
                }
                pre {
                    background: #f8f9fa;
                    padding: 15px;
                    border-radius: 5px;
                    overflow-x: auto;
                    border-left: 4px solid #3498db;
                }
                blockquote {
                    border-left: 4px solid #3498db;
                    margin: 1.5em 0;
                    padding-left: 20px;
                    color: #7f8c8d;
                    font-style: italic;
                }
                table {
                    border-collapse: collapse;
                    width: 100%;
                    margin: 1em 0;
                }
                th, td {
                    border: 1px solid #bdc3c7;
                    padding: 12px;
                    text-align: left;
                }
                th {
                    background: #ecf0f1;
                    font-weight: bold;
                }
                .document-meta {
                    background: #f8f9fa;
                    padding: 15px;
                    border-radius: 5px;
                    margin-bottom: 30px;
                    border-left: 4px solid #3498db;
                }
                .export-info {
                    margin-top: 50px;
                    padding-top: 20px;
                    border-top: 1px solid #bdc3c7;
                    font-size: 0.9em;
                    color: #7f8c8d;
                }
            </style>
            """

    def _convert_markdown_to_html(self) -> str:
        """Convert markdown content to HTML with XSS protection"""
        raw_html = markdown.markdown(
            self.document.markdown_content or '',
            extensions=['codehilite', 'fenced_code', 'tables', 'toc']
        )
        # SECURITY: Sanitize HTML to prevent XSS attacks
        return bleach.clean(
            raw_html,
            tags=ALLOWED_TAGS,
            attributes=ALLOWED_ATTRIBUTES,
            strip=True
        )

    def _generate_document_metadata_html(self) -> str:
        """Generate HTML for document metadata section"""
        # Escape user-controlled content to prevent XSS
        title_escaped = html.escape(self.document.title or '')
        author_escaped = html.escape(self.document.author or '') if self.document.author else ''
        tags_escaped = ", ".join(html.escape(tag) for tag in self.document.get_tag_names())

        author_html = f'<p><strong>Author:</strong> {author_escaped}</p>' if author_escaped else ''
        tags_html = f'<p><strong>Tags:</strong> {tags_escaped}</p>' if tags_escaped else ''

        return f"""
        <div class="document-meta">
            <h1>{title_escaped}</h1>
            {author_html}
            <p><strong>Created:</strong> {self.document.created_at.strftime('%B %d, %Y at %I:%M %p')}</p>
            <p><strong>Last Updated:</strong> {self.document.updated_at.strftime('%B %d, %Y at %I:%M %p')}</p>
            {tags_html}
        </div>
        """

    def _generate_export_info_html(self) -> str:
        """Generate HTML for export information footer"""
        export_time = datetime.now(timezone.utc).strftime('%B %d, %Y at %I:%M %p UTC')
        return f"""
        <div class="export-info">
            <p>Exported from Minky Document Management System on {export_time}</p>
        </div>
        """

    def _build_html_template(self, css_styles: str, meta_html: str, content_html: str, export_info: str) -> str:
        """Build complete HTML document from components"""
        title_escaped = html.escape(self.document.title or '')
        # SECURITY: Add Content-Security-Policy to prevent XSS
        return f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <meta http-equiv="Content-Security-Policy" content="default-src 'self'; script-src 'none'; style-src 'unsafe-inline'; img-src 'self' data: https:;">
            <title>{title_escaped}</title>
            {css_styles}
        </head>
        <body>
            {meta_html}
            {content_html}
            {export_info}
        </body>
        </html>
        """

    def export_to_html(self, include_styles=True):
        """Export document to HTML format"""
        html_content = self._convert_markdown_to_html()
        css_styles = self._generate_css_styles() if include_styles else ""
        meta_html = self._generate_document_metadata_html()
        export_info = self._generate_export_info_html()

        return self._build_html_template(css_styles, meta_html, html_content, export_info)
    
    def export_to_pdf(self):
        """Export document to PDF format using WeasyPrint"""
        raise NotImplementedError("PDF export is currently disabled due to WeasyPrint dependencies")
        # html_content = self.export_to_html(include_styles=True)
        
        # Additional CSS for PDF
        # pdf_css = CSS(string="""
        #     @page {
        #         margin: 1in;
        #         @bottom-center {
        #             content: counter(page) " of " counter(pages);
        #             font-size: 10px;
        #             color: #666;
        #         }
        #     }
        #     body { font-size: 12px; }
        #     h1 { page-break-before: avoid; }
        #     pre, table { page-break-inside: avoid; }
        # """)
        
        # pdf_file = os.path.join(self.temp_dir, f"{self.document.id}_{_sanitize_title_for_filename(self.document.title)}.pdf")
        # HTML(string=html_content).write_pdf(pdf_file, stylesheets=[pdf_css])
        
        # return pdf_file
    
    def export_to_docx(self):
        """Export document to DOCX format"""
        doc = DocxDocument()

        # Add title
        self._add_docx_title(doc)

        # Add metadata table
        self._add_docx_metadata(doc)

        doc.add_page_break()

        # Convert markdown to paragraphs
        self._add_docx_content(doc)

        # Add footer
        self._add_docx_footer(doc)

        # Save document
        docx_file = os.path.join(self.temp_dir, f"{self.document.id}_{_sanitize_title_for_filename(self.document.title)}.docx")
        doc.save(docx_file)

        return docx_file

    def _add_docx_title(self, doc: DocxDocument) -> None:
        """Add title to DOCX document"""
        title = doc.add_heading(self.document.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_docx_metadata(self, doc: DocxDocument) -> None:
        """Add metadata table to DOCX document"""
        meta_table = doc.add_table(rows=0, cols=2)
        meta_table.style = 'Table Grid'

        if self.document.author:
            row = meta_table.add_row()
            row.cells[0].text = 'Author'
            row.cells[1].text = self.document.author

        row = meta_table.add_row()
        row.cells[0].text = 'Created'
        row.cells[1].text = self.document.created_at.strftime('%B %d, %Y at %I:%M %p')

        row = meta_table.add_row()
        row.cells[0].text = 'Last Updated'
        row.cells[1].text = self.document.updated_at.strftime('%B %d, %Y at %I:%M %p')

        if self.document.get_tag_names():
            row = meta_table.add_row()
            row.cells[0].text = 'Tags'
            row.cells[1].text = ', '.join(self.document.get_tag_names())

    def _add_docx_content(self, doc: DocxDocument) -> None:
        """Convert markdown to DOCX paragraphs"""
        lines = self.document.markdown_content.split('\n')

        # SECURITY: Limit line count to prevent DoS
        if len(lines) > MAX_EXPORT_LINES:
            raise ValueError(f"Document has too many lines for export: {len(lines)} exceeds {MAX_EXPORT_LINES}")

        current_paragraph = ""

        for line in lines:
            current_paragraph = self._process_docx_line(doc, line.strip(), current_paragraph)

        if current_paragraph:
            doc.add_paragraph(current_paragraph)

    def _process_docx_line(self, doc: DocxDocument, line: str, current_paragraph: str) -> str:
        """Process a single markdown line for DOCX conversion"""
        if line.startswith('# '):
            return self._add_docx_heading(doc, line[2:], 1, current_paragraph)

        if line.startswith('## '):
            return self._add_docx_heading(doc, line[3:], 2, current_paragraph)

        if line.startswith('### '):
            return self._add_docx_heading(doc, line[4:], 3, current_paragraph)

        if line.startswith('```'):
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
            return ""

        if line == '':
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
            return ""

        return (current_paragraph + " " + line) if current_paragraph else line

    def _add_docx_heading(self, doc: DocxDocument, text: str, level: int, current_paragraph: str) -> str:
        """Add heading to DOCX document and return empty paragraph"""
        if current_paragraph:
            doc.add_paragraph(current_paragraph)
        doc.add_heading(text, level=level)
        return ""

    def _add_docx_footer(self, doc: DocxDocument) -> None:
        """Add footer to DOCX document"""
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"Exported from Minky on {datetime.now(timezone.utc).strftime('%B %d, %Y')}")
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def export_to_markdown(self):
        """Export document as clean markdown with metadata"""
        import yaml

        # SECURITY: Use yaml.safe_dump to properly escape special characters
        # This prevents YAML injection attacks via document title/author/tags
        metadata_dict = {
            'title': self.document.title or 'Untitled',
            'author': self.document.author or 'Unknown',
            'created': self.document.created_at.isoformat() if self.document.created_at else None,
            'updated': self.document.updated_at.isoformat() if self.document.updated_at else None,
            'tags': list(self.document.get_tag_names()),
            'exported': datetime.now(timezone.utc).isoformat()
        }

        yaml_content = yaml.safe_dump(metadata_dict, default_flow_style=False, allow_unicode=True)
        metadata = f"---\n{yaml_content}---\n\n"
        
        content = metadata + self.document.markdown_content
        
        md_file = os.path.join(self.temp_dir, f"{self.document.id}_{_sanitize_title_for_filename(self.document.title)}.md")
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return md_file
    
    def _get_export_safe_dict(self) -> dict:
        """SECURITY: Get document data safe for export (excludes internal IDs and owner PII)"""
        return {
            'title': self.document.title,
            'author': self.document.author,
            'created_at': self.document.created_at.isoformat() if self.document.created_at else None,
            'updated_at': self.document.updated_at.isoformat() if self.document.updated_at else None,
            'markdown_content': self.document.markdown_content,
            'is_public': self.document.is_public,
            'tags': list(self.document.get_tag_names()),
            # SECURITY: Exclude user_id, owner object, category internal details, document_metadata
        }

    def export_to_json(self):
        """Export document as JSON with safe metadata (excludes internal IDs)"""
        data = {
            # SECURITY: Use export-safe dict instead of full to_dict()
            'document': self._get_export_safe_dict(),
            'export_info': {
                'exported_at': datetime.now(timezone.utc).isoformat(),
                'format': 'json',
                'version': '1.0'
            }
        }
        
        json_file = os.path.join(self.temp_dir, f"{self.document.id}_{_sanitize_title_for_filename(self.document.title)}.json")
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        return json_file
    
    def export_bundle(self, formats=['html', 'pdf', 'docx', 'markdown', 'json']):
        """Export document in multiple formats as a ZIP bundle"""
        zip_file = os.path.join(self.temp_dir, f"{self.document.id}_{_sanitize_title_for_filename(self.document.title)}_bundle.zip")

        with zipfile.ZipFile(zip_file, 'w', zipfile.ZIP_DEFLATED) as zf:
            for format_type in formats:
                self._add_format_to_bundle(zf, format_type)

        return zip_file

    def _add_format_to_bundle(self, zf: zipfile.ZipFile, format_type: str) -> None:
        """Add a single format export to the bundle"""
        try:
            self._export_format_to_zip(zf, format_type)
        except Exception as e:
            # SECURITY: Log detailed error server-side, provide generic message in export
            logger.error(f"Export format {format_type} failed: {e}")
            zf.writestr(f"export_errors_{format_type}.txt",
                       f"Export to {format_type} format failed. Please contact support.")

    def _export_format_to_zip(self, zf: zipfile.ZipFile, format_type: str) -> None:
        """Export a specific format and add to ZIP"""
        title_truncated = _sanitize_title_for_filename(self.document.title)

        if format_type == 'html':
            content = self.export_to_html()
            zf.writestr(f"{title_truncated}.html", content.encode('utf-8'))
        elif format_type == 'pdf':
            pdf_path = self.export_to_pdf()
            zf.write(pdf_path, f"{title_truncated}.pdf")
        elif format_type == 'docx':
            docx_path = self.export_to_docx()
            zf.write(docx_path, f"{title_truncated}.docx")
        elif format_type == 'markdown':
            md_path = self.export_to_markdown()
            zf.write(md_path, f"{title_truncated}.md")
        elif format_type == 'json':
            json_path = self.export_to_json()
            zf.write(json_path, f"{title_truncated}.json")

def export_document(document, format_type='html', **options):
    """Convenience function to export a document"""
    exporter = DocumentExporter(document)

    try:
        result = _export_by_format(exporter, format_type, options)
        if format_type in ['html', 'json']:
            # For string returns, cleanup immediately
            exporter.cleanup()
        return result
    except Exception:
        # SECURITY: Always cleanup on exception to prevent temp file accumulation
        exporter.cleanup()
        raise
    # For file returns (pdf, docx, etc.), cleanup is handled by the route after sending


def _export_by_format(exporter: DocumentExporter, format_type: str, options: dict):
    """Export document using the specified format"""
    format_handlers = {
        'html': lambda: exporter.export_to_html(**options),
        'pdf': lambda: exporter.export_to_pdf(),
        'docx': lambda: exporter.export_to_docx(),
        'markdown': lambda: exporter.export_to_markdown(),
        'json': lambda: exporter.export_to_json(),
        'bundle': lambda: exporter.export_bundle(**options),
    }

    handler = format_handlers.get(format_type)
    if not handler:
        raise ValueError(f"Unsupported format: {format_type}")

    return handler()